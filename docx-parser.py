from pprint import pprint
from docx import Document

doc = Document('TC.docx')
table_start = 0
table_end = 2

def get_table(tb):
    table = {}
    table['name'] = doc.tables[tb].cell(0, 0).text + ' ' + doc.tables[tb].cell(0, 1).text
    table['description'] = doc.tables[tb].cell(1, 2).text
    table['requirements'] = doc.tables[tb].cell(3, 2).text

    step_num = 0
    table['steps'] = []
    while True:
        row = 5 + step_num
        try:
            step_desc = doc.tables[tb].cell(row, 0).text
            step_result = doc.tables[tb].cell(row, 2).text
        except IndexError:
            # print("End of test steps")
            break
        table['steps'].append((step_desc, step_result))
        step_num += 1

    return table

for tb in range(table_start, table_end):
    table = get_table(tb)
    print("Table ID:",tb)
    pprint(table)